"""Convierte Documentacion_BuenaBurger_POS.md a .docx usando el mismo convertidor del proyecto."""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INPUT  = r"C:\Users\PabloCantero\Desktop\TFG\buenaburger-pos-v1\docs\Documentacion_BuenaBurger_POS.md"
OUTPUT = r"C:\Users\PabloCantero\Desktop\TFG\buenaburger-pos-v1\docs\Documentacion_BuenaBurger_POS.docx"

doc = Document()

style_normal = doc.styles['Normal']
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(11)

for h_name, h_size, bold in [
    ('Heading 1', 16, True),
    ('Heading 2', 14, True),
    ('Heading 3', 13, True),
    ('Heading 4', 12, True),
]:
    s = doc.styles[h_name]
    s.font.name = 'Calibri'
    s.font.size = Pt(h_size)
    s.font.bold = bold
    s.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)


def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


INLINE_PAT = r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)'


def add_inline_runs(p, text):
    parts = re.split(INLINE_PAT, text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2]); run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1])
            run.font.name = 'Courier New'; run.font.size = Pt(10)
        elif part.startswith('*') and part.endswith('*'):
            run = p.add_run(part[1:-1]); run.italic = True
        else:
            p.add_run(part)


def add_inline_to_cell(cell, text):
    p = cell.paragraphs[0]
    add_inline_runs(p, text)


def add_table_from_md(lines):
    rows = []
    for line in lines:
        # saltar separador
        stripped = line.strip().strip('|')
        if re.match(r'^[\s\-:|]+$', stripped):
            continue
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        rows.append(cells)
    if not rows:
        return
    max_cols = max(len(r) for r in rows)
    rows = [r + [''] * (max_cols - len(r)) for r in rows]
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = 'Table Grid'
    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = ''
            add_inline_to_cell(cell, cell_text)
            cell.paragraphs[0].paragraph_format.space_after  = Pt(2)
            cell.paragraphs[0].paragraph_format.space_before = Pt(2)
            if i == 0:
                set_cell_bg(cell, '1F497D')
                for run in cell.paragraphs[0].runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            elif i % 2 == 0:
                set_cell_bg(cell, 'E8EEF7')


def clean(text):
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'\*(.*?)\*',     r'\1', text)
    text = re.sub(r'`(.*?)`',       r'\1', text)
    return text.strip()


def add_paragraph_with_inline(text, style=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    add_inline_runs(p, text)
    return p


with open(INPUT, encoding='utf-8') as f:
    lines = f.readlines()

i = 0
while i < len(lines):
    line = lines[i].rstrip('\n')

    # Bloque de código
    if line.strip().startswith('```'):
        i += 1
        code_lines = []
        while i < len(lines) and not lines[i].strip().startswith('```'):
            code_lines.append(lines[i].rstrip('\n'))
            i += 1
        code_text = '\n'.join(code_lines)
        p = doc.add_paragraph()
        run = p.add_run(code_text)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        p.paragraph_format.left_indent = Inches(0.3)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'),   'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'),  'F2F2F2')
        p._p.get_or_add_pPr().append(shading)
        i += 1
        continue

    # Tabla
    if line.strip().startswith('|'):
        table_lines = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            table_lines.append(lines[i].rstrip('\n'))
            i += 1
        add_table_from_md(table_lines)
        doc.add_paragraph()
        continue

    # Encabezados
    m = re.match(r'^(#{1,4})\s+(.*)', line)
    if m:
        level = len(m.group(1))
        text  = m.group(2)
        hmap  = {1: 'Heading 1', 2: 'Heading 2', 3: 'Heading 3', 4: 'Heading 4'}
        doc.add_paragraph(clean(text), style=hmap.get(level, 'Heading 4'))
        i += 1
        continue

    # Blockquote
    if line.strip().startswith('>'):
        text = line.strip().lstrip('> ').strip()
        p = doc.add_paragraph()
        run = p.add_run(clean(text))
        run.italic = True
        p.paragraph_format.left_indent = Inches(0.3)
        i += 1
        continue

    # Lista
    m_list = re.match(r'^(\s*)[-*]\s+(.*)', line)
    if m_list:
        indent = len(m_list.group(1))
        text   = m_list.group(2)
        p = doc.add_paragraph(style='List Bullet')
        add_inline_runs(p, text)
        if indent > 0:
            p.paragraph_format.left_indent = Inches(0.3 * (indent // 2 + 1))
        i += 1
        continue

    # Separador horizontal
    if re.match(r'^---+\s*$', line):
        p = doc.add_paragraph()
        pPr   = p._p.get_or_add_pPr()
        pBdr  = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'),   'single')
        bottom.set(qn('w:sz'),    '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'AAAAAA')
        pBdr.append(bottom)
        pPr.append(pBdr)
        i += 1
        continue

    # Línea vacía
    if line.strip() == '':
        i += 1
        continue

    # Párrafo normal
    add_paragraph_with_inline(line.strip())
    i += 1

doc.save(OUTPUT)
print(f"OK: {OUTPUT}")
